import re
import os
from datetime import date as _date
from io import BytesIO
from typing import Optional
from fastapi import APIRouter, Depends, HTTPException, Query, UploadFile, File
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from pydantic import BaseModel
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from database import get_db
import models
from dotenv import load_dotenv
from openai import OpenAI
import requests
from bs4 import BeautifulSoup

load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))

router = APIRouter(prefix="/api/athento", tags=["athento"])

# Columnas del Excel de Athento (base 0)
COL_FECHA         = 0
COL_ABOGADO       = 1
COL_CAUSAL        = 2
COL_SUBCAUSAL     = 3
COL_FICHERO       = 4
COL_EXP           = 5
COL_ESTADO        = 6
# COL_COMPANIA     = 7
COL_RESUMEN       = 8
# COL_REGIONAL     = 9
COL_RADICADO      = 10
COL_JUZGADO       = 11
COL_ACCIONANTE    = 12

# ─────────────────────────────────────────────────────────────────────────────
# MAPA_SUBCAUSAL
# Mapea texto del campo "Motivo causal IA" / "Subcausal IA" del Excel de Athento
# al nombre del modelo de plantilla en nuestro sistema.
# Se evalúa en orden: primero subcausal (más específico), luego causal.
# Todos los textos deben estar en minúsculas sin espacios extra.
# ─────────────────────────────────────────────────────────────────────────────
MAPA_SUBCAUSAL = {

    # ── ACCESO A CITAS ────────────────────────────────────────────────────────
    # Formato Athento AI
    "eps_salud_acceso a citas":               "Acceso a Citas",
    "eps_salud_acceso de citas":              "Acceso a Citas",
    "acceso a citas":                         "Acceso a Citas",
    "acceso de citas":                        "Acceso a Citas",
    "acceso a cita":                          "Acceso a Citas",
    "acceso de cita":                         "Acceso a Citas",
    # Subcausales Athento → Acceso a Citas
    "programación de consultas":              "Acceso a Citas",
    "programacion de consultas":              "Acceso a Citas",
    "programación de procedimientos qx":      "Acceso a Citas",
    "programacion de procedimientos qx":      "Acceso a Citas",
    "programación de ayudas dx":              "Acceso a Citas",
    "programacion de ayudas dx":              "Acceso a Citas",
    "programación de procedimientos":         "Acceso a Citas",
    "programacion de procedimientos":         "Acceso a Citas",
    "demora en citas":                        "Acceso a Citas",
    "demora citas":                           "Acceso a Citas",
    "asignacion de citas":                    "Acceso a Citas",
    "asignación de citas":                    "Acceso a Citas",
    "citas medicas":                          "Acceso a Citas",
    "citas médicas":                          "Acceso a Citas",

    # ── TRANSPORTE ────────────────────────────────────────────────────────────
    "eps_salud_transporte":                   "Transporte",
    "eps_salud_transporte y viaticos":        "Transporte y Viaticos",
    "eps_salud_transporte y viáticos":        "Transporte y Viaticos",
    "transporte y viaticos":                  "Transporte y Viaticos",
    "transporte y viáticos":                  "Transporte y Viaticos",
    "transporte":                             "Transporte",
    "transporte intermunicipal":              "Transporte",
    "traslado paciente":                      "Transporte",

    # ── NO INVIMA / MEDICAMENTO IMPORTADO ────────────────────────────────────
    "eps_salud_medicamentos":                 "Medicamento Importado",  # fallback
    "no indicación invima":                   "No INVIMA",
    "no indicacion invima":                   "No INVIMA",
    "sin registro invima":                    "No INVIMA",
    "sin registro sanitario":                 "No INVIMA",
    "medicamento sin invima":                 "No INVIMA",
    "medicamento importado":                  "Medicamento Importado",
    "medicamento no pbs":                     "Medicamento Importado",
    "medicamentos agotados/desabastecido por laboratorio": "Medicamento Importado",
    "medicamentos agotados":                  "Medicamento Importado",
    "medicamento no disponible en farmacia":  "Medicamento Importado",
    "medicamento agotado":                    "Medicamento Importado",
    "desabastecimiento":                      "Medicamento Importado",

    # ── SILLA DE RUEDAS ───────────────────────────────────────────────────────
    "eps_salud_silla de ruedas":              "Silla de Ruedas",
    "silla de ruedas":                        "Silla de Ruedas",
    "silla de ruedas electrica":              "Silla de Ruedas",
    "silla de ruedas eléctrica":              "Silla de Ruedas",
    "dispositivo ortopédico":                 "Silla de Ruedas",

    # ── NO PBS / PRESUPUESTOS MÁXIMOS ─────────────────────────────────────────
    "eps_salud_servicios especiales":         "No PBS / Presupuestos Maximos",
    "eps_salud_no pbs":                       "No PBS / Presupuestos Maximos",
    "no pbs":                                 "No PBS / Presupuestos Maximos",
    "presupuesto maximo":                     "No PBS / Presupuestos Maximos",
    "presupuesto máximo":                     "No PBS / Presupuestos Maximos",
    "presupuestos maximos":                   "No PBS / Presupuestos Maximos",
    "presupuestos máximos":                   "No PBS / Presupuestos Maximos",
    "tecnologia no pbs":                      "No PBS / Presupuestos Maximos",
    "tecnología no pbs":                      "No PBS / Presupuestos Maximos",
    "servicio de enfermería":                 "No PBS / Presupuestos Maximos",
    "servicio de enfermeria":                 "No PBS / Presupuestos Maximos",
    "nutriciones":                            "No PBS / Presupuestos Maximos",
    "nutrición enteral":                      "No PBS / Presupuestos Maximos",
    "nutricion enteral":                      "No PBS / Presupuestos Maximos",

    # ── CARENCIA ACTUAL DE OBJETO ─────────────────────────────────────────────
    "eps_salud-novedades administrativas":    "Carencia Actual de Objeto",
    "eps_salud_novedades administrativas":    "Carencia Actual de Objeto",
    "carencia actual de objeto":              "Carencia Actual de Objeto",
    "carencia de objeto":                     "Carencia Actual de Objeto",
    "hecho sobreviniente":                    "Carencia Actual de Objeto",
    "cambios de ips":                         "Carencia Actual de Objeto",
    "cambio de ips":                          "Carencia Actual de Objeto",
    "cambio ips":                             "Carencia Actual de Objeto",
    "traslado ips":                           "Carencia Actual de Objeto",

    # ── FALTA DE LEGITIMACIÓN EN LA CAUSA ────────────────────────────────────
    "falta de legitimacion en la causa":      "Falta de Legitimacion en la Causa",
    "falta de legitimación en la causa":      "Falta de Legitimacion en la Causa",
    "falta de legitimacion":                  "Falta de Legitimacion en la Causa",
    "falta de legitimación":                  "Falta de Legitimacion en la Causa",

    # ── TRATAMIENTO INTEGRAL ──────────────────────────────────────────────────
    "tratamiento integral":                   "Tratamiento Integral",
    "eps_salud_tratamiento integral":         "Tratamiento Integral",

    # ── CUIDADOR PRIMARIO ─────────────────────────────────────────────────────
    "cuidador primario":                      "Cuidador Primario",
    "cuidador":                               "Cuidador Primario",

    # ── AUTORIZACIONES ────────────────────────────────────────────────────────
    "eps_salud_autorizaciones":               "Autorizaciones",
    "autorizaciones":                         "Autorizaciones",
    "autorizaciones incompletas":             "Autorizaciones",
    "autorizaciones incompletas e incorrectas": "Autorizaciones",
    "adelanto en los tiempos de respuesta":   "Autorizaciones",

    # ── ATENCIÓN POR SOAT ─────────────────────────────────────────────────────
    "eps_salud_atencion por soat":            "Atencion por SOAT",
    "atencion por soat":                      "Atencion por SOAT",
}

ESTADO_SIN_CONCEPTO = "estructurar respuesta tutela"


def _extraer_partes_fichero(fichero: str) -> dict:
    """
    Extrae expediente y cédula del nombre del fichero.
    Formato: EXP-T-26-000023547 - 2026-0739 - 98765512
    """
    if not fichero:
        return {}
    partes = [p.strip() for p in str(fichero).split(" - ")]
    return {
        "expediente": partes[0] if len(partes) > 0 else "",
        "cedula":     partes[2] if len(partes) > 2 else "",
    }


def _clasificar(subcausal: str, causal: str, estado: str) -> dict:
    subcausal_l = (subcausal or "").lower().strip()
    causal_l    = (causal    or "").lower().strip()
    estado_l    = (estado    or "").lower().strip()

    # Buscar primero en subcausal, luego en causal si no hay match
    modelo = MAPA_SUBCAUSAL.get(subcausal_l) or MAPA_SUBCAUSAL.get(causal_l)
    sin_concepto = bool(modelo) or (estado_l == ESTADO_SIN_CONCEPTO)

    return {
        "sin_concepto":    sin_concepto,
        "modelo_sugerido": modelo or "",
        "auto_generable":  bool(modelo),
    }


def _extraer_info_resumen(resumen: str) -> dict:
    """Extrae campos estructurados del Resumen Juridico."""
    info = {"medida_provisional": "", "decision_juez": "", "hechos_resumen": ""}
    if not resumen:
        return info

    m = re.search(r'solicita medida provisional\?:\s*(S[IÍ]|NO)', resumen, re.IGNORECASE)
    if m:
        info["medida_provisional"] = m.group(1).upper()

    m = re.search(r'Decisi[oó]n del Juez:\s*(.+?)(?:HECHOS:|$)', resumen, re.IGNORECASE | re.DOTALL)
    if m:
        info["decision_juez"] = m.group(1).strip()[:300]

    m = re.search(r'HECHOS:\s*(.+?)(?:PRETENSIONES:|RESUMEN M[EÉ]DICO:|$)', resumen, re.IGNORECASE | re.DOTALL)
    if m:
        info["hechos_resumen"] = m.group(1).strip()[:400]

    return info


@router.post("/procesar")
async def procesar_excel(archivo: UploadFile = File(...), db: Session = Depends(get_db)):
    if not archivo.filename.endswith((".xlsx", ".xls")):
        raise HTTPException(status_code=400, detail="Solo se aceptan archivos .xlsx")

    contenido = await archivo.read()
    if not contenido:
        raise HTTPException(status_code=400, detail="El archivo está vacío.")

    try:
        wb = openpyxl.load_workbook(BytesIO(contenido), read_only=True, data_only=True)
        ws = wb.active
    except Exception as e:
        raise HTTPException(status_code=400, detail=f"No se pudo leer el Excel: {e}")

    tutelas = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        # Asegurar que la fila tiene al menos los primeros 13 elementos
        if len(row) < 7 or not any(v is not None for v in list(row)[:7]):
            continue

        # Acceder de forma segura con verificación de índice
        fichero    = str(row[COL_FICHERO]    or "") if len(row) > COL_FICHERO    else ""
        causal     = str(row[COL_CAUSAL]     or "") if len(row) > COL_CAUSAL     else ""
        subcausal  = str(row[COL_SUBCAUSAL]  or "") if len(row) > COL_SUBCAUSAL  else ""
        estado     = str(row[COL_ESTADO]     or "") if len(row) > COL_ESTADO     else ""
        abogado    = str(row[COL_ABOGADO]    or "") if len(row) > COL_ABOGADO    else ""
        resumen    = str(row[COL_RESUMEN]    or "") if len(row) > COL_RESUMEN    else ""
        fecha_raw  = row[COL_FECHA] if len(row) > COL_FECHA else None

        # Campos directos del Excel ampliado
        radicado   = str(row[COL_RADICADO]   or "") if len(row) > COL_RADICADO   else ""
        juzgado    = str(row[COL_JUZGADO]    or "") if len(row) > COL_JUZGADO    else ""
        accionante = str(row[COL_ACCIONANTE] or "") if len(row) > COL_ACCIONANTE else ""

        # Normalizar juzgado (quitar doble espacio)
        juzgado = " ".join(juzgado.split())

        # Cédula y expediente vienen del nombre del fichero
        partes = _extraer_partes_fichero(fichero)
        cedula = partes.get("cedula", "")
        expediente = partes.get("expediente", fichero)

        clasif = _clasificar(subcausal, causal, estado)
        extra  = _extraer_info_resumen(resumen)

        fecha_str = ""
        if fecha_raw:
            try:
                fecha_str = str(fecha_raw)[:10]
            except Exception:
                fecha_str = str(fecha_raw)

        # Si accionante no viene en columna directa, intentar extraerlo del resumen
        if not accionante:
            m = re.search(r'interpuesta por ([A-ZÁÉÍÓÚÑ][A-ZÁÉÍÓÚÑ\s]+?) contra', resumen)
            if m:
                accionante = m.group(1).strip()

        tutelas.append({
            "expediente":        expediente,
            "radicado":          radicado,
            "cedula":            cedula,
            "accionante":        accionante,
            "juzgado":           juzgado,
            "fecha_vencimiento": fecha_str,
            "abogado":           abogado,
            "causal":            causal,
            "subcausal":         subcausal,
            "estado":            estado,
            "sin_concepto":      clasif["sin_concepto"],
            "modelo_sugerido":   clasif["modelo_sugerido"],
            "auto_generable":    clasif["auto_generable"],
            "medida_provisional": extra["medida_provisional"],
            "decision_juez":     extra["decision_juez"],
            "hechos_resumen":    extra["hechos_resumen"],
            "resumen":           resumen[:300] + ("..." if len(resumen) > 300 else ""),
        })

    wb.close()

    # ── Guardar causales detectados en BD para uso del Athento directo ─────
    guardados = 0
    for t in tutelas:
        if t["modelo_sugerido"] and t["expediente"]:
            exp = t["expediente"].strip()
            existente = db.query(models.AthentoMotivoCausal).filter_by(
                uuid_athento=f"exp:{exp}"
            ).first()
            if existente:
                existente.modelo    = t["modelo_sugerido"]
                existente.subcausal = t["subcausal"] or ""
            else:
                db.add(models.AthentoMotivoCausal(
                    uuid_athento  = f"exp:{exp}",
                    expediente    = exp,
                    modelo        = t["modelo_sugerido"],
                    subcausal     = t["subcausal"] or "",
                    asignado_por  = f"Excel AI · {t['abogado'][:30]}",
                    notas         = f"causal={t['causal']} | estado={t['estado']}",
                ))
                guardados += 1
    db.commit()

    total = len(tutelas)
    sin_concepto_list = [t for t in tutelas if t["sin_concepto"]]
    por_modelo = {}
    for t in tutelas:
        m_key = t["modelo_sugerido"] or t["causal"] or "Sin clasificar"
        por_modelo[m_key] = por_modelo.get(m_key, 0) + 1

    return {
        "total":        total,
        "sin_concepto": len(sin_concepto_list),
        "con_concepto": total - len(sin_concepto_list),
        "por_modelo":   por_modelo,
        "tutelas":      tutelas,
    }


@router.get("/resumen")
async def resumen_athento(
    compania: str = Query("EPS", description="EPS | ARL"),
):
    """
    Resumen ejecutivo de la bandeja completa para el dashboard.
    Devuelve conteos reales por estado, abogado y vencimiento.
    """
    from datetime import date as _date
    try:
        from athento_client import fetch_bandeja
    except ImportError as e:
        raise HTTPException(status_code=500, detail=str(e))

    # Fetch TODAS (pendientes + contestadas) para tener el cuadro completo
    todas, error = fetch_bandeja(fecha=None, abogado=None, solo_pendientes=False, compania=compania)
    if error not in ("ok", "sin_datos"):
        raise HTTPException(status_code=502, detail=f"Athento: {error}")

    hoy  = _date.today().isoformat()           # "2026-06-07"
    circ = 2 * 3.14159 * 38                    # perímetro SVG donut

    # --- clasificar cada tutela ---
    completadas    = [t for t in todas if t.get("fecha_contestacion")]
    pendientes     = [t for t in todas if not t.get("fecha_contestacion")]
    sin_concepto   = [t for t in pendientes if t["sin_concepto"]]
    con_concepto   = [t for t in pendientes if not t["sin_concepto"]]

    # Vencidas: fecha_vencimiento < hoy (solo entre pendientes)
    def _vencida(t):
        v = (t.get("fecha_vencimiento") or "")[:10]
        return bool(v) and v < hoy

    def _para_hoy(t):
        v = (t.get("fecha_vencimiento") or "")[:10]
        return v == hoy

    vencidas   = [t for t in pendientes if _vencida(t)]
    para_hoy   = [t for t in pendientes if _para_hoy(t)]

    # --- por abogado ---
    from collections import defaultdict
    agrupad: dict[str, dict] = defaultdict(lambda: {
        "total": 0, "sin_concepto": 0, "vencidas": 0, "para_hoy": 0, "completadas": 0
    })
    for t in todas:
        a = t.get("abogado", "Sin asignar")
        agrupad[a]["total"] += 1
        if t.get("fecha_contestacion"):
            agrupad[a]["completadas"] += 1
        else:
            if t["sin_concepto"]:  agrupad[a]["sin_concepto"] += 1
            if _vencida(t):        agrupad[a]["vencidas"] += 1
            if _para_hoy(t):       agrupad[a]["para_hoy"] += 1

    por_abogado = sorted(
        [{"abogado": k, **v} for k, v in agrupad.items()],
        key=lambda x: x["total"], reverse=True
    )

    return {
        # conteos para el donut
        "completadas":   len(completadas),
        "con_concepto":  len(con_concepto),
        "sin_concepto":  len(sin_concepto),
        "vencidas":      len(vencidas),
        "para_hoy":      len(para_hoy),
        "total_activas": len(pendientes),
        "total_todas":   len(todas),
        # listas de tutelas urgentes
        "tutelas_hoy":     para_hoy[:20],
        "tutelas_vencidas": vencidas[:20],
        # desglose por abogado
        "por_abogado": por_abogado,
    }


@router.get("/bandeja")
async def bandeja_athento(
    abogado:  Optional[str] = Query(None,  description="Fragmento del nombre del abogado"),
    fecha:    Optional[str] = Query(None,  description="YYYY-MM-DD"),
    todos:    bool           = Query(False, description="Si true, incluye tutelas ya contestadas"),
    compania: str            = Query("EPS", description="EPS | ARL"),
):
    """
    Obtiene las tutelas directamente de Athento sin exportar Excel.
    Requiere ATHENTO_TOKEN o ATHENTO_SESSIONID en .env.
    """
    try:
        from athento_client import fetch_bandeja
    except ImportError as e:
        raise HTTPException(status_code=500, detail=f"athento_client no disponible: {e}")

    tutelas, error = fetch_bandeja(
        fecha=fecha or None,
        abogado=abogado or None,
        compania=compania,
        solo_pendientes=not todos,
    )

    if error == "session_expired":
        raise HTTPException(
            status_code=401,
            detail="La sesión de Athento expiró. Actualiza ATHENTO_SESSIONID o ATHENTO_TOKEN en .env.",
        )
    if error == "no_queue_id":
        raise HTTPException(status_code=500, detail="ATHENTO_QUEUE_ID no configurado en .env")
    if error not in ("ok", "sin_datos"):
        raise HTTPException(status_code=502, detail=f"Error Athento: {error}")

    sin_concepto_list = [t for t in tutelas if t["sin_concepto"]]
    auto_list         = [t for t in tutelas if t["auto_generable"]]

    # Agrupar por estado para mostrar distribución (equivalente a por_modelo del Excel)
    por_modelo: dict[str, int] = {}
    for t in tutelas:
        key = t["modelo_sugerido"] or t["estado"] or "Sin clasificar"
        por_modelo[key] = por_modelo.get(key, 0) + 1

    # Limpiar fecha_vencimiento: solo fecha "YYYY-MM-DD" sin hora
    from datetime import date as _date_cls
    hoy = _date_cls.today().isoformat()
    for t in tutelas:
        for campo in ("fecha_vencimiento", "fecha_notificacion", "fecha_expediente", "fecha_contestacion"):
            v = t.get(campo, "")
            if v and " " in v:
                t[campo] = v[:10]

        # ═══ CALCULAR RIESGO DE CONDENA ═══
        db = next(get_db())
        score_riesgo = 0

        # Score por juzgado
        juzgado_record = db.query(models.JuzgadoRiesgo).filter_by(juzgado=t.get("juzgado", "")).first()
        score_riesgo += juzgado_record.score if juzgado_record else 2

        # Score por modelo
        modelo_record = db.query(models.ModeloRiesgo).filter_by(modelo=t.get("modelo_sugerido", "")).first()
        score_riesgo += modelo_record.score if modelo_record else 2

        # Score por antigüedad
        fecha_exp = (t.get("fecha_expediente") or "")[:10]
        if fecha_exp:
            dias_antiguo = (_date_cls.fromisoformat(hoy) - _date_cls.fromisoformat(fecha_exp)).days
            if dias_antiguo < 3:
                score_riesgo += 1
            elif dias_antiguo < 7:
                score_riesgo += 2
            else:
                score_riesgo += 3
        else:
            score_riesgo += 2  # default

        # Determinar color y riesgo
        if score_riesgo <= 6:
            t["riesgo_color"] = "verde"
            t["riesgo_nivel"] = "Bajo"
        elif score_riesgo <= 9:
            t["riesgo_color"] = "amarillo"
            t["riesgo_nivel"] = "Medio"
        else:
            t["riesgo_color"] = "rojo"
            t["riesgo_nivel"] = "Alto"

        t["riesgo_score"] = score_riesgo
        db.close()

    return {
        "total":          len(tutelas),
        "sin_concepto":   len(sin_concepto_list),
        "auto_generable": len(auto_list),
        "con_concepto":   len(tutelas) - len(sin_concepto_list),
        "por_modelo":     por_modelo,
        "tutelas":        tutelas,
        "fuente":         "athento_directo",
    }


@router.get("/exportar")
async def exportar_excel(
    abogado:  Optional[str] = Query(None),
    fecha:    Optional[str] = Query(None),
    compania: str            = Query("EPS"),
):
    """Exporta la bandeja filtrada a .xlsx listo para descargar."""
    try:
        from athento_client import fetch_bandeja
    except ImportError as e:
        raise HTTPException(status_code=500, detail=str(e))

    tutelas, error = fetch_bandeja(
        fecha=fecha or None,
        abogado=abogado or None,
        solo_pendientes=True,
        compania=compania,
    )
    if error not in ("ok", "sin_datos"):
        raise HTTPException(status_code=502, detail=f"Athento: {error}")

    # ── Crear workbook ──────────────────────────────────────────────
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = f"Bandeja {compania}"

    # Estilos
    azul   = "1E3A8A"
    verde  = "166534"
    rojo   = "991B1B"
    ambar  = "92400E"
    gris   = "F8FAFC"

    hdr_fill  = PatternFill("solid", fgColor=azul)
    hdr_font  = Font(color="FFFFFF", bold=True, size=10)
    hdr_align = Alignment(horizontal="center", vertical="center", wrap_text=True)
    border    = Border(
        bottom=Side(style="thin", color="CBD5E1"),
        right=Side(style="thin",  color="E2E8F0"),
    )

    COLS = [
        ("Expediente",        18), ("Radicado",       13), ("Cédula",          13),
        ("Accionante",        26), ("Juzgado",         32), ("Abogado",         22),
        ("Regional",          12), ("Estado",          20), ("Medida Prov.",    10),
        ("Fecha Notif.",      14), ("Vencimiento",     14), ("Horas",            7),
        ("Sin Concepto",      10),
    ]

    # Cabecera
    for col_idx, (nombre, ancho) in enumerate(COLS, 1):
        cell = ws.cell(row=1, column=col_idx, value=nombre)
        cell.fill  = hdr_fill
        cell.font  = hdr_font
        cell.alignment = hdr_align
        cell.border = border
        ws.column_dimensions[get_column_letter(col_idx)].width = ancho
    ws.row_dimensions[1].height = 24
    ws.freeze_panes = "A2"

    hoy = _date.today().isoformat()

    # Filas
    for row_idx, t in enumerate(tutelas, 2):
        venc = (t.get("fecha_vencimiento") or "")[:10]
        vencida = bool(venc) and venc < hoy
        sin_c   = t.get("sin_concepto", False)

        # Color de fondo por urgencia
        if vencida:
            fill = PatternFill("solid", fgColor="FEE2E2")
        elif sin_c:
            fill = PatternFill("solid", fgColor="F0FDF4")
        elif venc and venc == hoy:
            fill = PatternFill("solid", fgColor="FFFBEB")
        else:
            fill = PatternFill("solid", fgColor="FFFFFF" if row_idx % 2 else gris)

        valores = [
            t.get("expediente",""), t.get("radicado",""),   t.get("cedula",""),
            t.get("accionante",""), t.get("juzgado",""),    t.get("abogado",""),
            t.get("regional",""),   t.get("estado",""),     t.get("medida_provisional",""),
            t.get("fecha_notificacion","")[:10] if t.get("fecha_notificacion") else "",
            venc,                   t.get("horas_otorgadas",""),
            "Sí" if sin_c else "No",
        ]
        for col_idx, val in enumerate(valores, 1):
            cell = ws.cell(row=row_idx, column=col_idx, value=val)
            cell.fill   = fill
            cell.border = border
            cell.alignment = Alignment(vertical="center")
            cell.font   = Font(size=10)
            # Vencidas en rojo
            if col_idx == 11 and vencida:
                cell.font = Font(size=10, bold=True, color=rojo)
            if col_idx == 13 and sin_c:
                cell.font = Font(size=10, bold=True, color=verde)

    # Metadatos y autofiltro
    ws.auto_filter.ref = f"A1:{get_column_letter(len(COLS))}{len(tutelas)+1}"
    ws.sheet_view.showGridLines = False

    nombre_archivo = f"Bandeja_{compania}_{_date.today().isoformat()}.xlsx"
    if abogado:
        nombre_archivo = f"Bandeja_{compania}_{abogado.split()[0]}_{_date.today().isoformat()}.xlsx"

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": f'attachment; filename="{nombre_archivo}"'},
    )


# ══════════════════════════════════════════════════════════════════════
#  CAUSALES — asignación manual por abogado
# ══════════════════════════════════════════════════════════════════════

class CausalIn(BaseModel):
    uuid_athento:  str
    expediente:    str = ""
    modelo:        str
    subcausal:     str = ""
    asignado_por:  str = "Abogado"
    notas:         str = ""


@router.post("/causales", status_code=201)
def guardar_causal(body: CausalIn, db: Session = Depends(get_db)):
    """Guarda o actualiza el motivo causal de una tutela de Athento."""
    existente = db.query(models.AthentoMotivoCausal).filter_by(
        uuid_athento=body.uuid_athento
    ).first()
    if existente:
        existente.modelo       = body.modelo
        existente.subcausal    = body.subcausal
        existente.asignado_por = body.asignado_por
        existente.notas        = body.notas
        existente.expediente   = body.expediente or existente.expediente
    else:
        existente = models.AthentoMotivoCausal(
            uuid_athento  = body.uuid_athento,
            expediente    = body.expediente,
            modelo        = body.modelo,
            subcausal     = body.subcausal,
            asignado_por  = body.asignado_por,
            notas         = body.notas,
        )
        db.add(existente)
    db.commit()
    db.refresh(existente)
    return {"ok": True, "id": existente.id, "modelo": existente.modelo}


@router.get("/causales")
def listar_causales(db: Session = Depends(get_db)):
    """Devuelve todos los causales asignados como dict {uuid: modelo}."""
    rows = db.query(models.AthentoMotivoCausal).all()
    return {
        r.uuid_athento: {
            "modelo":       r.modelo,
            "subcausal":    r.subcausal or "",
            "asignado_por": r.asignado_por or "",
            "notas":        r.notas or "",
        }
        for r in rows
    }


@router.delete("/causales/{uuid_athento}")
def borrar_causal(uuid_athento: str, db: Session = Depends(get_db)):
    db.query(models.AthentoMotivoCausal).filter_by(uuid_athento=uuid_athento).delete()
    db.commit()
    return {"ok": True}


# ─── CREDENCIALES DINÁMICAS ─────────────────────────────────────────────────────
_athento_creds = {}  # Almacena credenciales en memoria (persisten en esta sesión del servidor)


class AthentoCredsIn(BaseModel):
    sessionid: str
    csrftoken: str


@router.post("/update-creds")
def actualizar_credenciales(body: AthentoCredsIn):
    """Actualiza credenciales de Athento en tiempo real sin reinicar servidor."""
    global _athento_creds
    _athento_creds = {"sessionid": body.sessionid, "csrftoken": body.csrftoken}
    # También actualizar athento_client
    import athento_client
    athento_client.SESSIONID = body.sessionid
    athento_client.CSRFTOKEN = body.csrftoken
    return {"ok": True, "message": "Credenciales actualizadas"}


@router.get("/creds-status")
def estado_credenciales():
    """Verifica si hay credenciales válidas."""
    import athento_client
    tiene_creds = bool(athento_client.API_TOKEN or (athento_client.SESSIONID and athento_client.CSRFTOKEN))
    return {
        "autenticado": tiene_creds,
        "tipo": "API_TOKEN" if athento_client.API_TOKEN else "SESSION",
        "sessionid_activo": bool(athento_client.SESSIONID),
        "csrftoken_activo": bool(athento_client.CSRFTOKEN),
    }


# ─── SINCRONIZACIÓN AUTOMÁTICA DE CAUSALES ──────────────────────────────────────
@router.post("/sync-causales")
def sincronizar_causales(
    compania: str = Query("EPS", description="EPS o ARL"),
    abogado: str | None = Query(None, description="Filtro opcional por nombre de abogado"),
    db: Session = Depends(get_db),
):
    """
    Sincroniza causales de Athento usando Selenium para extraer "Motivo causal IA" y "Subcausal IA".
    Procesa tutelas que no tengan causal asignado aún.
    """
    import athento_client
    from routers.athento import MAPA_SUBCAUSAL

    # Obtener bandeja sin filtro de causales
    tutelas, error = athento_client.fetch_bandeja(
        fecha=None,
        abogado=abogado,
        compania=compania,
        solo_pendientes=True,
    )

    if error != "ok" and error != "sin_datos":
        raise HTTPException(status_code=400, detail=f"Error fetching bandeja: {error}")

    # Obtener causales ya guardados
    existentes = db.query(models.AthentoMotivoCausal).all()
    uuids_guardados = {r.uuid_athento for r in existentes}

    procesados = 0
    guardados = 0
    errores = []

    for tutela in tutelas:
        uuid = tutela.get("uuid", "")
        expediente = tutela.get("expediente", "")

        # Saltar si ya tiene causal
        if uuid in uuids_guardados:
            continue

        procesados += 1

        # Extraer causal con Selenium
        print(f"[{procesados}] Extrayendo {uuid}...")
        resultado = athento_client._extraer_causal_documento(uuid)

        if not resultado:
            errores.append(f"{expediente}: no extrae causal")
            continue

        motivo = resultado.get("motivo_causal", "").strip().lower()
        subcausal = resultado.get("subcausal_ia", "").strip().lower()

        # Mapear a modelo interno
        modelo = MAPA_SUBCAUSAL.get(motivo) or MAPA_SUBCAUSAL.get(subcausal) or ""

        if modelo:
            try:
                db.add(models.AthentoMotivoCausal(
                    uuid_athento=uuid,
                    expediente=expediente,
                    modelo=modelo,
                    subcausal=resultado.get("subcausal_ia", ""),  # guardar original
                    asignado_por=f"Auto-Selenium · {tutela.get('abogado', '')[:25]}",
                    notas=f"motivo_causal_ia={resultado.get('motivo_causal')}",
                ))
                db.commit()
                guardados += 1
                print(f"  ✓ Guardado: {modelo}")
            except Exception as e:
                errores.append(f"{expediente}: {str(e)}")
                print(f"  ✗ Error guardando: {e}")
        else:
            errores.append(f"{expediente}: no mapea a modelo (motivo={motivo})")

    return {
        "procesados": procesados,
        "guardados": guardados,
        "errores": errores,
        "total_procesados": len(tutelas),
    }


# ─── CÁLCULO DE RIESGO DE CONDENA ────────────────────────────────────────────────
@router.post("/calcular-riesgo")
def calcular_riesgo_tutela(
    juzgado: str,
    modelo: str,
    dias_antiguo: int = 0,
    db: Session = Depends(get_db),
):
    """
    Calcula el riesgo de condena para una tutela SURA basado en:
    - Juzgado: historicamente favorable/desfavorable a SURA
    - Modelo: tipo de tutela (medicamento, acceso, etc)
    - Antigüedad: cuántos días lleva sin responder

    Devuelve:
    - score: 4-6 (verde), 7-9 (amarillo), 10+ (rojo)
    - color: "verde" | "amarillo" | "rojo"
    - riesgo: "Bajo" | "Medio" | "Alto"
    """

    score = 0
    desglose = {}

    # 1. Score por juzgado (1-3)
    juzgado_record = db.query(models.JuzgadoRiesgo).filter_by(juzgado=juzgado).first()
    score_juzgado = juzgado_record.score if juzgado_record else 2  # default neutral
    score += score_juzgado
    desglose["juzgado"] = {"nombre": juzgado[:40], "score": score_juzgado, "notas": juzgado_record.notas if juzgado_record else ""}

    # 2. Score por modelo (1-3)
    modelo_record = db.query(models.ModeloRiesgo).filter_by(modelo=modelo).first()
    score_modelo = modelo_record.score if modelo_record else 2
    score += score_modelo
    desglose["modelo"] = {"nombre": modelo, "score": score_modelo, "notas": modelo_record.notas if modelo_record else ""}

    # 3. Score por antigüedad (1-3)
    # < 3 días → 1, 3-7 días → 2, > 7 días → 3
    if dias_antiguo < 3:
        score_antiguedad = 1
    elif dias_antiguo < 7:
        score_antiguedad = 2
    else:
        score_antiguedad = 3
    score += score_antiguedad
    desglose["antiguedad"] = {"dias": dias_antiguo, "score": score_antiguedad}

    # 4. Determinar color y nivel de riesgo
    if score <= 6:
        color = "verde"
        riesgo = "Bajo"
        confianza = 70  # % probabilidad de que SURA gane
    elif score <= 9:
        color = "amarillo"
        riesgo = "Medio"
        confianza = 50
    else:
        color = "rojo"
        riesgo = "Alto"
        confianza = 30

    return {
        "score": score,
        "color": color,
        "riesgo": riesgo,
        "confianza_ganar": confianza,
        "desglose": desglose,
        "recomendacion": f"{'⚠️ REVISAR CON EXPERTO' if color == 'rojo' else 'Proceder normalmente'}"
    }


# Endpoint para inicializar scores por defecto (admin)
@router.post("/init-riesgos")
def inicializar_scores_riesgo(db: Session = Depends(get_db)):
    """Inicializa la tabla de riesgos con valores por defecto para SURA."""

    # Limpiar
    db.query(models.JuzgadoRiesgo).delete()
    db.query(models.ModeloRiesgo).delete()

    # Modelos de riesgo (valores iniciales basados en experiencia SURA)
    modelos_riesgo = [
        ("Medicamento Importado", 1, "SURA suele ganar estos casos"),
        ("Transporte y Viaticos", 1, "Bajo riesgo, normalmente resuelto"),
        ("Acceso a Citas", 3, "SURA tiende a perder, requiere estrategia especial"),
        ("No PBS / Presupuestos Maximos", 2, "Riesgo medio, depende de jurisprudencia"),
        ("Silla de Ruedas", 2, "Caso por caso"),
        ("Autorizaciones", 2, "Riesgo medio"),
        ("Atencion por SOAT", 1, "Bajo riesgo"),
        ("Carencia Actual de Objeto", 1, "Bajo riesgo"),
        ("Tratamiento Integral", 3, "Alto riesgo, requiere experto"),
        ("Cuidador Primario", 3, "Alto riesgo"),
    ]

    for modelo, score, notas in modelos_riesgo:
        db.add(models.ModeloRiesgo(
            modelo=modelo,
            score=score,
            notas=notas,
            actualizado_por="Sistema"
        ))

    db.commit()
    return {"inicializado": True, "modelos": len(modelos_riesgo), "juzgados": 0}


# ─── CONCEPTOS PRE-ESTABLECIDOS ──────────────────────────────────────────────
@router.get("/conceptos/{modelo}")
def obtener_conceptos_modelo(modelo: str, db: Session = Depends(get_db)):
    """Obtiene los conceptos pre-establecidos disponibles para un modelo."""
    conceptos = db.query(models.ConceptoPreestablecido).filter(
        models.ConceptoPreestablecido.modelo == modelo,
        models.ConceptoPreestablecido.activo == True
    ).order_by(models.ConceptoPreestablecido.orden).all()

    if not conceptos:
        return {"disponibles": False, "conceptos": []}

    return {
        "disponibles": True,
        "modelo": modelo,
        "conceptos": [
            {
                "id": c.id,
                "nombre": c.nombre,
                "variante": c.variante,
                "texto_concepto": c.texto_concepto,
            }
            for c in conceptos
        ]
    }


@router.get("/conceptos")
def listar_todos_conceptos(db: Session = Depends(get_db)):
    """Lista todos los conceptos disponibles agrupados por modelo."""
    conceptos = db.query(models.ConceptoPreestablecido).filter(
        models.ConceptoPreestablecido.activo == True
    ).all()

    # Agrupar por modelo
    por_modelo = {}
    for c in conceptos:
        if c.modelo not in por_modelo:
            por_modelo[c.modelo] = []
        por_modelo[c.modelo].append({
            "id": c.id,
            "nombre": c.nombre,
            "variante": c.variante,
            "largo": len(c.texto_concepto),
        })

    return {
        "total_modelos": len(por_modelo),
        "total_conceptos": len(conceptos),
        "por_modelo": por_modelo,
    }


# ─── GENERACIÓN DE PLANTILLAS ───────────────────────────────────────────────
def _reemplazar_en_docx(plantilla_bytes: bytes, reemplazos: dict) -> bytes:
    """
    Reemplaza placeholders en DOCX usando python-docx (seguro, no daña XML).
    """
    from docx import Document
    import io

    doc = Document(io.BytesIO(plantilla_bytes))

    # Reemplazar en párrafos
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            for marcador, valor in reemplazos.items():
                if marcador in run.text:
                    run.text = run.text.replace(marcador, str(valor) if valor else marcador)

    # Reemplazar en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        for marcador, valor in reemplazos.items():
                            if marcador in run.text:
                                run.text = run.text.replace(marcador, str(valor) if valor else marcador)

    # Guardar a bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()


class GenerarPlantillaRequest(BaseModel):
    modelo: str
    concepto_texto: str
    radicado: str = ""
    accionante: str = ""
    juzgado: str = ""
    representante_legal: str = ""
    cedula_representante: str = ""


@router.post("/generar-plantilla-con-concepto")
def generar_plantilla_con_concepto(
    req: GenerarPlantillaRequest,
    db: Session = Depends(get_db),
):
    """
    Genera un documento ODT desde la plantilla principal reemplazando placeholders
    con el concepto técnico proporcionado en el request.
    """
    import os
    from datetime import datetime

    # DEBUG: Log de lo que se recibe
    print(f"DEBUG: Concepto texto length: {len(req.concepto_texto)}")
    print(f"DEBUG: Concepto primeros 100 chars: {req.concepto_texto[:100]}")
    print(f"DEBUG: Accionante: {req.accionante}")
    print(f"DEBUG: Radicado: {req.radicado}")

    # 2. Cargar plantilla principal desde plantillas/
    plantilla_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "plantillas", "Transporte.docx")

    if not os.path.exists(plantilla_path):
        raise HTTPException(status_code=500, detail=f"Plantilla no encontrada")

    try:
        with open(plantilla_path, 'rb') as f:
            plantilla_contenido = f.read()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error leyendo plantilla: {str(e)}")

    # 3. Preparar reemplazos
    fecha_hoy = datetime.now().strftime("%d de %B de %Y").replace(
        "January", "enero").replace("February", "febrero").replace("March", "marzo").replace(
        "April", "abril").replace("May", "mayo").replace("June", "junio").replace(
        "July", "julio").replace("August", "agosto").replace("September", "septiembre").replace(
        "October", "octubre").replace("November", "noviembre").replace("December", "diciembre")

    reemplazos = {
        "[CONCEPTO TECNICO]": req.concepto_texto,
        "[ACCIONANTE]": req.accionante,
        "[RADICADO]": req.radicado,
        "[JUZGADO]": req.juzgado,
        "[REPRESENTANTE_LEGAL]": req.representante_legal,
        "[CEDULA_REPRESENTANTE]": req.cedula_representante,
        "[NUMERO_RS]": "",
        "[FECHA_CONTESTACION]": fecha_hoy,
    }

    print(f"DEBUG: Reemplazos preparados: {list(reemplazos.keys())}")

    # 4. Reemplazar en el ODT
    try:
        contenido_modificado = _reemplazar_en_docx(plantilla_contenido, reemplazos)
        print(f"DEBUG: ODT procesado exitosamente")
    except Exception as e:
        print(f"ERROR procesando ODT: {str(e)}")
        raise HTTPException(status_code=500, detail=f"Error procesando ODT: {str(e)}")

    # 5. Retornar documento generado
    nombre_archivo = f"Contestacion_{req.modelo}_{req.radicado or 'sin_radicado'}.docx"

    return StreamingResponse(
        iter([contenido_modificado]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={nombre_archivo}"}
    )


@router.post("/generar-plantilla/{modelo}")
def generar_plantilla(
    modelo: str,
    concepto_id: int | None = None,
    radicado: str = "",
    accionante: str = "",
    db: Session = Depends(get_db),
):
    """
    Genera un documento ODT desde la plantilla principal reemplazando placeholders
    con concepto técnico y otros datos.

    Placeholders reemplazados:
    - [CONCEPTO TECNICO]
    - [ACCIONANTE]
    - [RADICADO]
    - [JUZGADO]
    - [REPRESENTANTE_LEGAL]
    - [CEDULA_REPRESENTANTE]
    - [NUMERO_RS]
    - [FECHA_CONTESTACION]
    """
    import os
    from datetime import datetime

    # 1. Obtener concepto de BD
    if concepto_id:
        concepto = db.query(models.ConceptoPreestablecido).filter_by(
            id=concepto_id,
            activo=True
        ).first()
    else:
        concepto = db.query(models.ConceptoPreestablecido).filter(
            models.ConceptoPreestablecido.modelo == modelo,
            models.ConceptoPreestablecido.activo == True
        ).order_by(models.ConceptoPreestablecido.orden).first()

    if not concepto:
        raise HTTPException(status_code=404, detail=f"No hay concepto pre-establecido para {modelo}")

    # 2. Cargar plantilla principal desde plantillas/
    plantilla_path = os.path.join(os.path.dirname(os.path.dirname(__file__)), "plantillas", "Transporte.docx")

    if not os.path.exists(plantilla_path):
        raise HTTPException(status_code=500, detail=f"Plantilla no encontrada")

    try:
        with open(plantilla_path, 'rb') as f:
            plantilla_contenido = f.read()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error leyendo plantilla: {str(e)}")

    # 3. Preparar reemplazos
    fecha_hoy = datetime.now().strftime("%d de %B de %Y").replace(
        "January", "enero").replace("February", "febrero").replace("March", "marzo").replace(
        "April", "abril").replace("May", "mayo").replace("June", "junio").replace(
        "July", "julio").replace("August", "agosto").replace("September", "septiembre").replace(
        "October", "octubre").replace("November", "noviembre").replace("December", "diciembre")

    reemplazos = {
        "[CONCEPTO TECNICO]": concepto.texto_concepto,
        "[ACCIONANTE]": accionante,
        "[RADICADO]": radicado,
        "[JUZGADO]": "",  # Se podría agregar si viene en el request
        "[REPRESENTANTE_LEGAL]": "",
        "[CEDULA_REPRESENTANTE]": "",
        "[NUMERO_RS]": "",
        "[FECHA_CONTESTACION]": fecha_hoy,
    }

    # 4. Reemplazar en el ODT
    try:
        contenido_modificado = _reemplazar_en_docx(plantilla_contenido, reemplazos)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error procesando ODT: {str(e)}")

    # 5. Retornar documento generado
    nombre_archivo = f"Contestacion_{modelo}_{radicado or 'sin_radicado'}.docx"

    return StreamingResponse(
        iter([contenido_modificado]),
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": f"attachment; filename={nombre_archivo}"}
    )


# ─── AGENTE IA PARA SUGERIR MODELO ──────────────────────────────────────────
@router.post("/sugerir-modelo")
def sugerir_modelo(
    servicio_solicitado: str,
    resumen_hechos: str = "",
):
    """
    Usa Claude Haiku para sugerir el modelo de tutela más apropiado
    basado en el servicio solicitado y los hechos del caso.
    """
    try:
        modelos_disponibles = [
            "Transporte",
            "Transporte y Viaticos",
            "Medicamento Importado",
            "No INVIMA",
            "Silla de Ruedas",
            "Acceso a Citas",
            "Tratamiento Integral",
            "Cuidador Primario",
            "No PBS / Presupuestos Maximos",
        ]

        prompt = f"""Eres un experto en tutelas de salud. Basado en los hechos descritos,
sugiere el modelo de tutela más apropriado.

SERVICIO SOLICITADO: {servicio_solicitado}

HECHOS DEL CASO:
{resumen_hechos if resumen_hechos else "(sin detalles adicionales)"}

MODELOS DISPONIBLES:
{chr(10).join([f"- {m}" for m in modelos_disponibles])}

Responde SOLO con el nombre del modelo, sin explicación. Ej: "Medicamento Importado"
Si no estás seguro, elige el más probable."""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            max_tokens=100,
            messages=[{"role": "user", "content": prompt}]
        )

        modelo_sugerido = response.choices[0].message.content.strip()

        # Validar que el modelo sugerido esté en la lista
        if modelo_sugerido not in modelos_disponibles:
            # Si no coincide exactamente, buscar el más cercano
            modelo_sugerido = modelos_disponibles[0]

        return {
            "modelo_sugerido": modelo_sugerido,
            "confianza": "alta" if servicio_solicitado else "media",
            "razon": "Analizado por Claude IA"
        }

    except Exception as e:
        return {
            "error": str(e),
            "modelo_sugerido": None
        }


# ─── ADAPTACIÓN INTELIGENTE DE CONCEPTOS ────────────────────────────────────
class AdaptarConceptoRequest(BaseModel):
    modelo: str
    concepto_original: str
    hechos_caso: str = ""
    servicio_solicitado: str = ""


@router.post("/adaptar-concepto")
def adaptar_concepto(req: AdaptarConceptoRequest):
    """
    Adapta un concepto pre-establecido basándose en los hechos específicos del caso.
    Por ejemplo: si es Acceso a Citas, identifica si es cita general o especialista
    y personaliza el concepto en consecuencia.
    """
    try:
        prompt = f"""Eres un abogado experto en tutelas de salud. Tu tarea es adaptar y personalizar
un concepto técnico pre-establecido basándose en los hechos específicos de un caso.

MODELO: {req.modelo}
SERVICIO SOLICITADO: {req.servicio_solicitado}

HECHOS DEL CASO:
{req.hechos_caso if req.hechos_caso else "(sin detalles adicionales)"}

CONCEPTO PRE-ESTABLECIDO (base):
{req.concepto_original}

INSTRUCCIONES:
1. Analiza los hechos y el servicio solicitado
2. Identifica detalles clave (ej: si es cita general vs especialista, medicamento importado vs nacional, tipo de transporte, etc.)
3. Adapta el concepto pre-establecido para que sea específico a estos hechos
4. Mantén la estructura legal y el tono profesional
5. Haz que sea más persuasivo y personalizado al caso concreto
6. No agregues comillas ni explicaciones, solo el concepto adaptado

CONCEPTO ADAPTADO:"""

        response = client.chat.completions.create(
            model="gpt-4o-mini",
            max_tokens=2000,
            messages=[{"role": "user", "content": prompt}]
        )

        concepto_adaptado = response.choices[0].message.content.strip()

        return {
            "concepto_adaptado": concepto_adaptado,
            "modelo": req.modelo,
            "personalizado": True
        }

    except Exception as e:
        return {
            "error": str(e),
            "concepto_adaptado": None
        }


# ─── EXTRAER TUTELA POR UUID (SIMPLIFICADO SIN SELENIUM) ─────────────────────
@router.get("/extraer-tutela/{uuid}")
def extraer_tutela(uuid: str, db: Session = Depends(get_db)):
    """
    Extrae información de una tutela en Athento usando la API directamente.
    Usa Claude para procesar y extraer campos clave.
    """
    try:
        base_url = os.getenv("ATHENTO_BASE", "https://sura.athento.com")
        sessionid = os.getenv("ATHENTO_SESSIONID", "")
        csrftoken = os.getenv("ATHENTO_CSRFTOKEN", "")

        # Crear sesión de requests
        s = requests.Session()
        s.cookies.set("sessionid", sessionid)
        s.cookies.set("csrftoken", csrftoken)
        s.headers.update({"X-CSRFToken": csrftoken})

        # Obtener el documento
        url_doc = f"{base_url}/file/view/{uuid}"
        resp = s.get(url_doc, timeout=10)

        if resp.status_code != 200:
            return {
                "error": f"No se pudo acceder al documento (HTTP {resp.status_code})",
                "uuid": uuid
            }

        # Parsear HTML
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(resp.text, "html.parser")

        # Extraer todo el texto visible
        page_text = soup.get_text(separator="\n")[:2000]  # Primeros 2000 caracteres

        # Extraer campos de formulario
        datos = {
            "uuid": uuid,
            "accionante": "",
            "radicado": "",
            "juzgado": "",
            "servicio": "",
            "resumen_hechos": "",
            "modelo_sugerido": ""
        }

        # Buscar inputs y textareas (BeautifulSoup)
        try:
            for label in soup.find_all("label"):
                label_text = label.get_text().lower()
                parent = label.parent
                if parent:
                    input_field = parent.find("input") or parent.find("textarea") or parent.find("select")
                    if input_field:
                        valor = input_field.get("value", "") or input_field.get_text("").strip()
                        valor = str(valor).strip() if valor else ""

                        if "accionante" in label_text:
                            datos["accionante"] = valor[:100]
                        elif "radicado" in label_text:
                            datos["radicado"] = valor[:50]
                        elif "juzgado" in label_text:
                            datos["juzgado"] = valor[:100]
                        elif "servicio" in label_text or "medicamento" in label_text:
                            datos["servicio"] = valor[:150]
                        elif "hechos" in label_text or "descripción" in label_text:
                            datos["resumen_hechos"] = valor[:500]
        except Exception as e:
            print(f"Error parseando campos: {e}")

        # Usar Claude para detectar modelo basado en el contenido
        if not datos["modelo_sugerido"]:
            try:
                response = client.chat.completions.create(
                    model="gpt-4o-mini",
                    max_tokens=50,
                    messages=[{
                        "role": "user",
                        "content": f"""Basado en este texto del documento, ¿cuál es el modelo de tutela?

{page_text}

Responde SOLO con: Medicamento Importado | Transporte | Acceso a Citas | No INVIMA | Silla de Ruedas | Tratamiento Integral | Cuidador Primario | Otro
"""
                    }]
                )
                modelo = response.choices[0].message.content.strip()
                if "Medicamento" in modelo or "medicamento" in modelo:
                    datos["modelo_sugerido"] = "Medicamento Importado"
                elif "Transporte" in modelo or "transporte" in modelo:
                    datos["modelo_sugerido"] = "Transporte"
                elif "Acceso" in modelo or "cita" in modelo:
                    datos["modelo_sugerido"] = "Acceso a Citas"
                elif "INVIMA" in modelo:
                    datos["modelo_sugerido"] = "No INVIMA"
                elif "Silla" in modelo or "ruedas" in modelo:
                    datos["modelo_sugerido"] = "Silla de Ruedas"
                elif "Tratamiento" in modelo or "integral" in modelo:
                    datos["modelo_sugerido"] = "Tratamiento Integral"
                elif "Cuidador" in modelo:
                    datos["modelo_sugerido"] = "Cuidador Primario"
            except:
                pass

        return datos

    except Exception as e:
        print(f"Error en extraer_tutela: {e}")
        return {
            "error": f"Error: {str(e)}",
            "uuid": uuid
        }


# ─── DETECTAR TUTELA ABIERTA EN ATHENTO ──────────────────────────────────────
@router.get("/detectar-tutela-abierta")
def detectar_tutela_abierta(db: Session = Depends(get_db)):
    """
    Detecta qué tutela está abierta actualmente en Athento
    y extrae información del documento (accionante, radicado, hechos, etc.)
    """
    try:
        from selenium import webdriver
        from selenium.webdriver.common.by import By
        from selenium.webdriver.support.ui import WebDriverWait
        from selenium.webdriver.support import expected_conditions as EC
        from webdriver_manager.chrome import ChromeDriverManager
        from selenium.webdriver.chrome.service import Service
        import time
        import re

        base_url = os.getenv("ATHENTO_BASE", "https://sura.athento.com")
        sessionid = os.getenv("ATHENTO_SESSIONID", "")
        csrftoken = os.getenv("ATHENTO_CSRFTOKEN", "")

        # Configurar Chrome
        options = webdriver.ChromeOptions()
        options.add_argument("--headless")
        options.add_argument("--no-sandbox")
        options.add_argument("--disable-dev-shm-usage")
        options.add_argument("--window-size=1920,1080")

        driver = webdriver.Chrome(
            service=Service(ChromeDriverManager().install()),
            options=options
        )

        # Navegar a Athento y agregar cookies
        driver.get(f"{base_url}/")
        time.sleep(1)
        driver.add_cookie({"name": "sessionid", "value": sessionid})
        driver.add_cookie({"name": "csrftoken", "value": csrftoken})

        # Navegar a la cola para obtener el último documento abierto
        driver.get(f"{base_url}/queues/api/queue/content_queue/")
        time.sleep(2)

        # Intentar obtener el UUID de la URL actual si está en un documento
        current_url = driver.current_url
        uuid_match = re.search(r'/file/view/([a-f0-9-]+)', current_url)

        if not uuid_match:
            # Buscar en el HTML para obtener enlaces a documentos
            try:
                links = driver.find_elements(By.TAG_NAME, "a")
                for link in links[:1]:  # Primer documento
                    href = link.get_attribute("href")
                    if "/file/view/" in href:
                        uuid_match = re.search(r'/file/view/([a-f0-9-]+)', href)
                        if uuid_match:
                            break
            except:
                pass

        if not uuid_match:
            driver.quit()
            return {
                "error": "No se encontró tutela abierta. Abre una en Athento primero.",
                "uuid": None
            }

        uuid = uuid_match.group(1)

        # Navegar al documento
        driver.get(f"{base_url}/file/view/{uuid}")
        time.sleep(2)

        # Extraer información del formulario
        datos = {
            "uuid": uuid,
            "accionante": "",
            "radicado": "",
            "juzgado": "",
            "servicio": "",
            "resumen_hechos": "",
            "modelo_sugerido": ""
        }

        try:
            # Buscar campos en el documento
            inputs = driver.find_elements(By.TAG_NAME, "input")
            selects = driver.find_elements(By.TAG_NAME, "select")
            textareas = driver.find_elements(By.TAG_NAME, "textarea")

            # Intentar extraer valores
            for elem in inputs + selects + textareas:
                try:
                    label = elem.find_element(By.XPATH, "./preceding-sibling::label").text.lower()
                    value = elem.get_attribute("value") or elem.text

                    if "accionante" in label and not datos["accionante"]:
                        datos["accionante"] = value.strip()
                    elif "radicado" in label and not datos["radicado"]:
                        datos["radicado"] = value.strip()
                    elif "juzgado" in label and not datos["juzgado"]:
                        datos["juzgado"] = value.strip()
                    elif ("servicio" in label or "medicamento" in label) and not datos["servicio"]:
                        datos["servicio"] = value.strip()
                    elif "hechos" in label and not datos["resumen_hechos"]:
                        datos["resumen_hechos"] = value.strip()
                except:
                    pass

            driver.quit()

            # Si no tiene servicio, intentar obtenerlo del radicado o hechos
            if not datos["servicio"] and datos["resumen_hechos"]:
                # Usar IA para sugerir servicio de los hechos
                palabras_clave = {
                    "medicamento": "Medicamento",
                    "transporte": "Transporte",
                    "cita": "Acceso a Citas",
                    "especialista": "Acceso a Citas",
                    "silla": "Silla de Ruedas",
                    "psicólogo": "Tratamiento Integral"
                }
                for palabra, sugerencia in palabras_clave.items():
                    if palabra.lower() in datos["resumen_hechos"].lower():
                        datos["modelo_sugerido"] = sugerencia
                        break

            return datos

        except Exception as e:
            driver.quit()
            print(f"Error extrayendo datos: {e}")
            return {
                "error": f"Error extrayendo información: {str(e)}",
                "uuid": uuid
            }

    except Exception as e:
        print(f"Error detectando tutela: {e}")
        return {
            "error": f"No se pudo conectar a Athento: {str(e)}",
            "uuid": None
        }

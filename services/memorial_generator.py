"""
Generador de memoriales judiciales (Cumplimiento, Impugnación, Incidente de Desacato).
- Cumplimiento: usa la plantilla real Memorial_Cumplimiento.docx (reemplazo de marcadores).
- Impugnación / Incidente de Desacato: generados con python-docx desde código.
"""

import re
import zipfile
from datetime import date
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, RGBColor, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from services.word_generator import _procesar_xml, _insertar_advertencia, PLANTILLAS_DIR

BASE_DIR = Path(__file__).resolve().parent.parent
MEMORIALES_DIR = BASE_DIR / "memoriales"

TIPOS_MEMORIAL = ["Cumplimiento", "Impugnación", "Incidente de Desacato"]

# Nombre del archivo de plantilla para cada tipo basado en plantilla real
PLANTILLA_MEMORIAL = {
    "Cumplimiento": "Memorial_Cumplimiento.docx",
}


# ─────────────────────────────────────────────
# Helpers de formato
# ─────────────────────────────────────────────

def _set_font(run, size_pt: int = 12, bold: bool = False, color: RGBColor = None):
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color


def _parrafo(doc: Document, texto: str, alineacion=WD_ALIGN_PARAGRAPH.JUSTIFY,
             size_pt: int = 12, bold: bool = False, espacio_antes: int = 0,
             espacio_despues: int = 6, color: RGBColor = None) -> None:
    p = doc.add_paragraph()
    p.alignment = alineacion
    p.paragraph_format.space_before = Pt(espacio_antes)
    p.paragraph_format.space_after = Pt(espacio_despues)
    p.paragraph_format.line_spacing = Pt(18)
    run = p.add_run(texto)
    _set_font(run, size_pt, bold, color)


def _titulo_seccion(doc: Document, texto: str) -> None:
    _parrafo(doc, texto, alineacion=WD_ALIGN_PARAGRAPH.LEFT,
             size_pt=12, bold=True, espacio_antes=10, espacio_despues=4)


def _advertencia_borrador(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    run = p.add_run(
        "*** BORRADOR — PARA REVISIÓN DEL ABOGADO RESPONSABLE — NO RADICAR ***"
    )
    _set_font(run, size_pt=10, bold=True, color=RGBColor(0xFF, 0x00, 0x00))


def _configurar_margenes(doc: Document) -> None:
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.0)


def _encabezado_radicacion(doc: Document, datos: dict) -> None:
    """Bloque superior: ciudad/fecha a la derecha, destinatario a la izquierda."""
    ciudad = datos.get("ciudad_departamento", "")
    fecha = datos.get("fecha_memorial", "")
    _parrafo(doc, f"{ciudad}, {fecha}",
             alineacion=WD_ALIGN_PARAGRAPH.RIGHT, espacio_despues=14)

    _parrafo(doc, f"Señor(a) Juez(a)",
             alineacion=WD_ALIGN_PARAGRAPH.LEFT, espacio_despues=2)
    _parrafo(doc, datos.get("juzgado", ""),
             alineacion=WD_ALIGN_PARAGRAPH.LEFT, espacio_despues=2)
    _parrafo(doc, "E.S.D.",
             alineacion=WD_ALIGN_PARAGRAPH.LEFT, espacio_despues=14)


def _referencia(doc: Document, datos: dict, asunto: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(14)
    r1 = p.add_run("Ref: ")
    _set_font(r1, bold=True)
    r2 = p.add_run(
        f"Acción de Tutela – Radicado No. {datos.get('radicado', '')}\n"
        f"      Accionante: {(datos.get('accionante', '') or '').upper()}\n"
        f"      {asunto}"
    )
    _set_font(r2)


def _parrafo_comparece(doc: Document, datos: dict) -> None:
    rep = datos.get("representante_legal", "")
    ced = datos.get("cedula_representante", "")
    cargo = datos.get("cargo_representante", "representante legal")
    _parrafo(
        doc,
        f"{rep}, identificado(a) con C.C. No. {ced}, actuando en calidad de "
        f"{cargo} de EPS y Medicina Prepagada Suramericana S.A. (EPS SURA), "
        f"respetuosamente me dirijo a su Despacho para:",
        espacio_despues=10,
    )


def _firma(doc: Document, datos: dict) -> None:
    _parrafo(doc, "Atentamente,",
             alineacion=WD_ALIGN_PARAGRAPH.LEFT, espacio_antes=20, espacio_despues=40)
    rep = datos.get("representante_legal", "")
    ced = datos.get("cedula_representante", "")
    regional = datos.get("regional", "")
    _parrafo(doc, rep, alineacion=WD_ALIGN_PARAGRAPH.LEFT,
             bold=True, espacio_despues=2)
    _parrafo(doc, f"C.C. No. {ced}",
             alineacion=WD_ALIGN_PARAGRAPH.LEFT, espacio_despues=2)
    _parrafo(doc, f"EPS SURA – Regional {regional}",
             alineacion=WD_ALIGN_PARAGRAPH.LEFT, espacio_despues=2)


# ─────────────────────────────────────────────
# Generadores por tipo
# ─────────────────────────────────────────────

def _generar_cumplimiento(datos: dict) -> Document:
    doc = Document()
    _configurar_margenes(doc)
    _advertencia_borrador(doc)
    _encabezado_radicacion(doc, datos)
    _referencia(doc, datos,
                f"EPS SURA – No. Interno: {datos.get('numero_rs', '')}")

    _parrafo_comparece(doc, datos)

    _titulo_seccion(doc, "I. OBJETO DEL PRESENTE MEMORIAL")
    _parrafo(
        doc,
        "Dar cumplimiento a la orden judicial impartida en la sentencia de tutela "
        f"proferida el {datos.get('fecha_orden_judicial', '')}, informando al Despacho "
        "sobre las gestiones realizadas en acatamiento del fallo.",
    )

    _titulo_seccion(doc, "II. HECHOS DEL CUMPLIMIENTO")
    _parrafo(doc, datos.get("descripcion_cumplimiento", ""))

    _titulo_seccion(doc, "III. FECHA DE CUMPLIMIENTO EFECTIVO")
    _parrafo(
        doc,
        f"La orden judicial fue acatada el día {datos.get('fecha_cumplimiento_efectivo', '')}, "
        "dentro del término establecido en el fallo.",
    )

    soportes = datos.get("soportes_adjuntos", "")
    if soportes:
        _titulo_seccion(doc, "IV. SOPORTES ADJUNTOS")
        _parrafo(doc, soportes)

    obs = datos.get("observaciones_memorial", "")
    if obs:
        _titulo_seccion(doc, "V. OBSERVACIONES")
        _parrafo(doc, obs)

    _parrafo(
        doc,
        "De esta manera, EPS SURA da cabal cumplimiento a la orden impartida por "
        "su Despacho y pone a su disposición los soportes que así lo acreditan.",
        espacio_antes=8,
    )
    _firma(doc, datos)
    return doc


def _generar_impugnacion(datos: dict) -> Document:
    doc = Document()
    _configurar_margenes(doc)
    _advertencia_borrador(doc)
    _encabezado_radicacion(doc, datos)
    _referencia(doc, datos,
                f"Impugnación del fallo – EPS SURA No. Interno: {datos.get('numero_rs', '')}")

    _parrafo_comparece(doc, datos)

    _titulo_seccion(doc, "I. OBJETO DE LA IMPUGNACIÓN")
    tipo_fallo = datos.get("tipo_fallo", "")
    _parrafo(
        doc,
        f"Impugnar el fallo de tutela de primera instancia proferido el "
        f"{datos.get('fecha_fallo_impugnado', '')}, mediante el cual se "
        f"{tipo_fallo} la acción de tutela interpuesta por el accionante, "
        "por las razones que a continuación se exponen.",
    )

    _titulo_seccion(doc, "II. ARGUMENTOS DE LA IMPUGNACIÓN")
    _parrafo(doc, datos.get("argumentos_impugnacion", ""))

    _titulo_seccion(doc, "III. PRETENSIÓN")
    pretension = datos.get("pretension_impugnacion", "")
    if pretension:
        _parrafo(doc, pretension)
    else:
        _parrafo(
            doc,
            "Revocar el fallo de primera instancia y, en su lugar, negar las "
            "pretensiones de la acción de tutela, por no encontrarse acreditada "
            "la vulneración de los derechos fundamentales invocados.",
        )

    obs = datos.get("observaciones_memorial", "")
    if obs:
        _titulo_seccion(doc, "IV. OBSERVACIONES")
        _parrafo(doc, obs)

    _parrafo(
        doc,
        "En mérito de lo expuesto, respetuosamente solicitamos al Despacho "
        "dar trámite a la presente impugnación conforme a lo dispuesto en el "
        "artículo 32 del Decreto 2591 de 1991.",
        espacio_antes=8,
    )
    _firma(doc, datos)
    return doc


def _generar_desacato(datos: dict) -> Document:
    doc = Document()
    _configurar_margenes(doc)
    _advertencia_borrador(doc)
    _encabezado_radicacion(doc, datos)
    _referencia(doc, datos,
                f"Respuesta Incidente de Desacato – EPS SURA No. Interno: {datos.get('numero_rs', '')}")

    _parrafo_comparece(doc, datos)

    _titulo_seccion(doc, "I. OBJETO")
    _parrafo(
        doc,
        f"Dar respuesta al incidente de desacato abierto mediante auto del "
        f"{datos.get('fecha_apertura_incidente', '')}, manifestando las razones "
        "por las cuales EPS SURA ha dado o está dando cumplimiento a la orden judicial.",
    )

    _titulo_seccion(doc, "II. RELACIÓN DE HECHOS Y ACCIONES REALIZADAS")
    _parrafo(doc, datos.get("descripcion_incidente", ""))

    _titulo_seccion(doc, "III. ACCIONES EJECUTADAS PARA DAR CUMPLIMIENTO")
    _parrafo(doc, datos.get("acciones_realizadas", ""))

    fecha_cumpl = datos.get("fecha_cumplimiento_efectivo", "")
    if fecha_cumpl:
        _titulo_seccion(doc, "IV. FECHA DE CUMPLIMIENTO")
        _parrafo(
            doc,
            f"EPS SURA ejecutó las acciones necesarias para el cumplimiento de la "
            f"orden judicial el día {fecha_cumpl}.",
        )

    justif = datos.get("justificacion_desacato", "")
    if justif:
        _titulo_seccion(doc, "V. JUSTIFICACIÓN")
        _parrafo(doc, justif)

    soportes = datos.get("soportes_adjuntos", "")
    if soportes:
        _titulo_seccion(doc, "VI. SOPORTES")
        _parrafo(doc, soportes)

    _parrafo(
        doc,
        "Con fundamento en lo anterior, respetuosamente solicitamos al Despacho "
        "declarar el cumplimiento de la orden judicial y archivar el incidente de "
        "desacato, dado que EPS SURA ha satisfecho la pretensión del accionante.",
        espacio_antes=8,
    )
    _firma(doc, datos)
    return doc


# ─────────────────────────────────────────────
# Generación por plantilla real (Cumplimiento)
# ─────────────────────────────────────────────

def _mapa_cumplimiento(datos: dict) -> dict:
    d = datos.get
    cedula = (d("cedula_representante", "") or "").replace("CC ", "").replace("C.C. ", "").strip()
    return {
        "[JUZGADO]":                   d("juzgado", ""),
        "[ACCIONANTE]":                (d("accionante", "") or "").upper(),
        "[RADICADO]":                  d("radicado", ""),
        "[REPRESENTANTE_LEGAL]":       d("representante_legal", ""),
        "[NUMERO_CEDULA]":             cedula,
        "[CEDULA_REPRESENTANTE]":      cedula,
        "[TARJETA_PROFESIONAL]":       d("tarjeta_profesional", ""),
        "[CARGO_REPRESENTANTE]":       d("cargo_representante", "apoderada judicial"),
        "[CARGO_REPRESENTANTE_FORMAL]": d("cargo_representante", "Apoderada Judicial"),
        "[FECHA_ORDEN_JUDICIAL]":      d("fecha_orden_judicial", ""),
        "[DESCRIPCION_ORDEN]":         d("descripcion_orden", ""),
        "[DESCRIPCION_ORDEN_CONT]":    d("descripcion_orden", ""),
        "[CUMPLIMIENTO_1]":            d("cumplimiento_1", ""),
        "[CUMPLIMIENTO_2]":            d("cumplimiento_2", ""),
        "[CUMPLIMIENTO_3]":            d("cumplimiento_3", ""),
        "[FECHA_MEMORIAL]":            d("fecha_memorial", ""),
    }


def _generar_desde_plantilla(tipo: str, datos: dict) -> bytes:
    nombre = PLANTILLA_MEMORIAL[tipo]
    ruta = PLANTILLAS_DIR / nombre
    if not ruta.exists():
        raise FileNotFoundError(
            f"Plantilla no encontrada: {ruta}. "
            "Verifique que el archivo esté en la carpeta plantillas/."
        )
    mapa = _mapa_cumplimiento(datos)
    zip_in = BytesIO(ruta.read_bytes())
    zip_out = BytesIO()
    with zipfile.ZipFile(zip_in, "r") as zin:
        with zipfile.ZipFile(zip_out, "w", zipfile.ZIP_DEFLATED) as zout:
            for item in zin.infolist():
                data = zin.read(item.filename)
                if item.filename.endswith(".xml"):
                    try:
                        xml, _ = _procesar_xml(data.decode("utf-8"), mapa)
                        if item.filename == "word/document.xml":
                            xml = _insertar_advertencia(xml)
                        data = xml.encode("utf-8")
                    except Exception:
                        pass
                zout.writestr(item, data)
    return zip_out.getvalue()


# ─────────────────────────────────────────────
# Función principal
# ─────────────────────────────────────────────

_GENERADORES_DOCX = {
    "Impugnación":            _generar_impugnacion,
    "Incidente de Desacato":  _generar_desacato,
}


def generar_memorial(datos: dict) -> tuple[str, str]:
    """
    Genera el .docx del memorial.
    Retorna (nombre_archivo, ruta_completa).
    """
    tipo = datos.get("tipo_memorial", "")
    if not tipo:
        raise ValueError("Tipo de memorial no especificado.")

    rs = re.sub(r"[^A-Za-z0-9]", "", datos.get("numero_rs", "SINUS"))
    accionante = re.sub(r"\s+", "_", (datos.get("accionante", "SIN") or "SIN").upper())[:20]
    accionante = re.sub(r"[^A-Za-z0-9_]", "", accionante)
    tipo_slug = (re.sub(r"\s+", "_", tipo)
                 .replace("ó", "o").replace("é", "e").replace("ú", "u"))
    fecha_hoy = date.today().strftime("%Y%m%d")
    nombre_archivo = f"Memorial_{tipo_slug}_{rs}_{accionante}_{fecha_hoy}.docx"

    MEMORIALES_DIR.mkdir(parents=True, exist_ok=True)
    ruta_salida = MEMORIALES_DIR / nombre_archivo

    if tipo in PLANTILLA_MEMORIAL:
        # Generación basada en plantilla .docx real
        contenido = _generar_desde_plantilla(tipo, datos)
        ruta_salida.write_bytes(contenido)
    else:
        # Generación con python-docx desde código
        generador = _GENERADORES_DOCX.get(tipo)
        if not generador:
            raise ValueError(f"Tipo de memorial no reconocido: '{tipo}'")
        doc = generador(datos)
        doc.save(str(ruta_salida))

    return nombre_archivo, str(ruta_salida)

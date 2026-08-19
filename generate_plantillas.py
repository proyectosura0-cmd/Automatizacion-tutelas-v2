"""
Genera documentos desde plantilla principal reemplazando [CONCEPTO TECNICO]
con el contenido de los conceptos específicos por modelo.
"""
import os
import shutil
from pathlib import Path
from docx import Document as DocxDocument
from odf import opendoc
from odf.opendocument import OpenDocumentText
from odf.text import P, Span, Tab
import xml.etree.ElementTree as ET

# Rutas
CONCEPTOS_DIR = r"C:\Users\alvar\OneDrive\Escritorio\automatizacion -sura\Causales-Modelos\Conceptos"
PLANTILLA_DIR = r"C:\Users\alvar\OneDrive\Escritorio\automatizacion -sura\Causales-Modelos\Modelo Principal"
OUTPUT_DIR = r"C:\Users\alvar\tutelas-sura\plantillas_generadas"

# Crear directorio de salida
os.makedirs(OUTPUT_DIR, exist_ok=True)

def extraer_contenido_docx(ruta_docx: str) -> str:
    """Extrae todo el texto de un archivo DOCX."""
    try:
        doc = DocxDocument(ruta_docx)
        contenido = []
        for para in doc.paragraphs:
            if para.text.strip():
                contenido.append(para.text)
        # También extraer de tablas
        for tabla in doc.tables:
            for fila in tabla.rows:
                for celda in fila.cells:
                    if celda.text.strip():
                        contenido.append(celda.text)
        return "\n".join(contenido)
    except Exception as e:
        print(f"Error leyendo {ruta_docx}: {e}")
        return ""


def leer_plantilla_odt(ruta_odt: str) -> bytes:
    """Lee el contenido binario de la plantilla ODT."""
    try:
        with open(ruta_odt, 'rb') as f:
            return f.read()
    except Exception as e:
        print(f"Error leyendo plantilla ODT: {e}")
        return b""


def reemplazar_en_odt(contenido_odt: bytes, buscar: str, reemplazar: str) -> bytes:
    """
    Reemplaza texto en archivo ODT (funciona buscando en el XML interno).
    ODT es un ZIP con archivos XML internos.
    """
    try:
        import zipfile
        import io

        # ODT es un ZIP
        with zipfile.ZipFile(io.BytesIO(contenido_odt), 'r') as zip_read:
            # Leer content.xml (donde está el contenido principal)
            content_xml = zip_read.read('content.xml').decode('utf-8')

            # Hacer el reemplazo
            # NOTA: buscar está entre > y <, así que buscamos el texto directo
            content_xml = content_xml.replace(buscar, reemplazar)

            # Crear nuevo ZIP con el contenido modificado
            output = io.BytesIO()
            with zipfile.ZipFile(output, 'w', zipfile.ZIP_DEFLATED) as zip_write:
                # Copiar todos los archivos excepto content.xml
                for item in zip_read.infolist():
                    if item.filename == 'content.xml':
                        zip_write.writestr('content.xml', content_xml.encode('utf-8'))
                    else:
                        zip_write.writestr(item, zip_read.read(item.filename))

            return output.getvalue()
    except Exception as e:
        print(f"Error reemplazando en ODT: {e}")
        return contenido_odt


def generar_documento(modelo: str, concepto_contenido: str, salida: str):
    """Genera un documento desde la plantilla reemplazando el concepto."""
    plantilla_path = os.path.join(PLANTILLA_DIR, "MODELO PRINCIPAL.odt")

    if not os.path.exists(plantilla_path):
        print(f"❌ Plantilla no encontrada: {plantilla_path}")
        return False

    # Leer plantilla
    contenido = leer_plantilla_odt(plantilla_path)

    # Reemplazar [CONCEPTO TECNICO] con el contenido
    contenido_modificado = reemplazar_en_odt(
        contenido,
        "[CONCEPTO TECNICO]",
        concepto_contenido
    )

    # Guardar resultado
    try:
        with open(salida, 'wb') as f:
            f.write(contenido_modificado)
        print(f"✅ Generado: {salida}")
        return True
    except Exception as e:
        print(f"❌ Error guardando {salida}: {e}")
        return False


def main():
    """Genera documentos para todos los modelos."""
    # Mapeo de archivos de conceptos
    archivos_conceptos = {
        "Acesso a citas.docx": "Acceso a Citas",
        "Medicamento importado.docx": "Medicamento Importado",
        "No INVIMA.docx": "No INVIMA",
        "Transporte.docx": "Transporte",
        "Transporte y viaticos.docx": "Transporte y Viáticos",
        "Tratamiento integral.docx": "Tratamiento Integral",
        "cuidador primario.docx": "Cuidador Primario",
        "silla de ruedas.docx": "Silla de Ruedas",
    }

    generados = 0
    errores = 0

    print("🔄 Generando plantillas desde template principal...\n")

    for archivo_concepto, nombre_modelo in archivos_conceptos.items():
        ruta_concepto = os.path.join(CONCEPTOS_DIR, archivo_concepto)

        if not os.path.exists(ruta_concepto):
            print(f"⚠️  No encontrado: {archivo_concepto}")
            errores += 1
            continue

        # Extraer contenido del concepto
        contenido_concepto = extraer_contenido_docx(ruta_concepto)
        if not contenido_concepto:
            print(f"⚠️  Concepto vacío: {archivo_concepto}")
            errores += 1
            continue

        # Generar documento
        salida = os.path.join(OUTPUT_DIR, f"{nombre_modelo}.odt")
        if generar_documento(nombre_modelo, contenido_concepto, salida):
            generados += 1
        else:
            errores += 1

    print(f"\n📊 Resumen:")
    print(f"   ✅ Generados: {generados}")
    print(f"   ❌ Errores: {errores}")
    print(f"   📁 Ubicación: {OUTPUT_DIR}")


if __name__ == "__main__":
    main()
